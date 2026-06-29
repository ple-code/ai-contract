"""合同导出服务。

- build_revised_contract：基于 docxtpl 模板渲染「采购合同修订稿」（保留模板排版/字体/签章位）
- build_review_report：代码生成「合同审查报告」（结构化报告，非合同本体）
- convert_to_pdf：调用 LibreOffice headless 把 docx 转 pdf
"""
import io
import os
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from datetime import datetime
from lxml import etree
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docxtpl import DocxTemplate
from sqlalchemy import select

from ..config import resolve_upload_path
from ..models.clause import Clause, ClauseReviewState
from ..models.contract import Contract, ContractType, ContractVersion
from ..models.review import Finding, Review

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "采购合同模板.docx"

RISK_LABEL = {"high": "高风险", "medium": "中风险", "mid": "中风险", "low": "低风险"}
DECISION_LABEL = {"accept": "接受", "reject": "拒绝"}
STANCE_LABEL = {"buyer": "甲方（采购方）", "seller": "乙方（供应方）", "neutral": "中立"}

def _format_legal_basis(legal_basis) -> str:
    """兼容 AI 入库字段 article / article_no 不统一。"""
    if not legal_basis:
        return ""
    parts = []
    for lb in legal_basis:
        if not isinstance(lb, dict):
            continue
        law = lb.get("law") or ""
        article = lb.get("article_no") or lb.get("article") or lb.get("articles") or ""
        label = f"{law} {article}".strip()
        if label:
            parts.append(label)
    return "、".join(parts)


_PARTY_RE = re.compile(r"(甲方|需方|买方|采购方)[：:]\s*([^\s，,。；;]+)")
_PARTY_B_RE = re.compile(r"(乙方|供方|卖方|供应方|供货方)[：:]\s*([^\s，,。；;]+)")


def _extract_party_names(clauses: list[Clause]) -> tuple[str, str]:
    """从条款文本里尽力抽取甲乙方名称，找不到则用占位符。"""
    blob = "\n".join((c.title + " " + c.text) for c in clauses[:8])
    a = ""
    b = ""
    m = _PARTY_RE.search(blob)
    if m:
        a = m.group(2).strip("（(】［[")
    m = _PARTY_B_RE.search(blob)
    if m:
        b = m.group(2).strip("）)】］]")
    return (
        a or "××××有限公司（采购方）",
        b or "××××有限公司（供应方）",
    )


def _sign_place(clauses: list[Clause]) -> str:
    blob = "\n".join(c.text for c in clauses[:10])
    m = re.search(r"(签订地点|签署地)[：:]\s*([^\s，,。；;]+)", blob)
    return m.group(2).strip() if m else "________"


# ---------------- 采购合同修订稿（原文件原地改，B 方案） ----------------

def _set_para_text(p_elem, text: str) -> None:
    """清空段落所有 run（保留 pPr 段落格式壳），写入单行文本。"""
    for child in list(p_elem):
        if child.tag != qn("w:pPr"):
            p_elem.remove(child)
    if not text:
        return
    r = etree.SubElement(p_elem, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    t.set(qn("xml:space"), "preserve")


def _replace_clause_blocks(blocks: list, new_text: str) -> None:
    """把 clause 对应的 body 元素原地替换为 new_text（干净稿，无修订标记）。

    - blocks[0]（首个段落）保留段落壳、写入 new_text 第一行；
    - blocks[1:] 其余元素从 body 删除；
    - new_text 多行 → 后续行复用 blocks[0] 的 pPr 格式插入新段；
    - 含表格（w:tbl）的 clause 由调用方跳过，避免破坏表格结构。
    """
    if not blocks:
        return
    first = blocks[0]
    parent = first.getparent()
    if parent is None:
        return
    lines = new_text.split("\n")
    _set_para_text(first, lines[0])
    for b in blocks[1:]:
        if b.getparent() is not None:
            parent.remove(b)
    insert_after = first
    for line in lines[1:]:
        new_p = deepcopy(first)
        _set_para_text(new_p, line)
        insert_after.addnext(new_p)
        insert_after = new_p


async def build_revised_contract(db, version_id: int) -> io.BytesIO:
    """导出修订稿：在用户上传的原 docx 上原地替换「已应用建议」的条款文本。

    与早期 docxtpl 套模板方案不同，此方案保留原合同的排版/表格/签章位，
    只把 ClauseReviewState.applied=True 的条款段落，替换为当前 clause.text
    （apply 时已写入 finding.suggestion）。导出为干净稿，无 track changes 痕迹。
    """
    from fastapi import HTTPException
    from .doc_parser import parse_docx

    ver = await db.get(ContractVersion, version_id)
    if not ver or not ver.file_uri:
        raise HTTPException(500, "版本无源文件路径，无法导出修订稿")
    contract = await db.get(Contract, ver.contract_id)

    db_clauses = list((await db.execute(
        select(Clause).where(Clause.version_id == version_id).order_by(Clause.id)
    )).scalars().all())
    db_map = {c.code: c for c in db_clauses}
    states_list = list((await db.execute(
        select(ClauseReviewState).where(ClauseReviewState.version_id == version_id)
    )).scalars().all())
    states_map = {s.clause_code: s for s in states_list}

    # 打开原文件并解析；blocks 指向此 doc 的 body 子元素，用于原地替换
    doc_path = resolve_upload_path(ver.file_uri)
    if not doc_path.is_file():
        raise HTTPException(500, f"版本源文件不存在：{ver.file_uri}")

    doc = Document(str(doc_path))
    parsed = parse_docx(str(doc_path), _doc=doc)

    replaced = 0
    skipped_table = 0
    for pc in parsed:
        st = states_map.get(pc.code)
        if not (st and st.applied and pc.code in db_map):
            continue
        # 含表格的条款整段替换会破坏表格结构，跳过（保留原样）
        if any(getattr(b, "tag", None) == qn("w:tbl") for b in pc.blocks):
            skipped_table += 1
            continue
        _replace_clause_blocks(pc.blocks, db_map[pc.code].text)
        replaced += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------- 审查报告（代码生成，结构化） ----------------

async def build_review_report(db, version_id: int) -> io.BytesIO:
    ver = await db.get(ContractVersion, version_id)
    contract = await db.get(Contract, ver.contract_id)
    ct = (await db.execute(select(ContractType).where(ContractType.code == contract.type_code))).scalar_one_or_none()

    clauses = list((await db.execute(
        select(Clause).where(Clause.version_id == version_id).order_by(Clause.id)
    )).scalars().all())

    review = (await db.execute(
        select(Review).where(Review.version_id == version_id).order_by(Review.id.desc())
    )).scalars().first()

    findings_list = []
    if review:
        findings_list = list((await db.execute(
            select(Finding).where(Finding.review_id == review.id).order_by(Finding.id)
        )).scalars().all())

    states_list = list((await db.execute(
        select(ClauseReviewState).where(ClauseReviewState.version_id == version_id)
    )).scalars().all())

    findings_map: dict[str, list] = {}
    for f in findings_list:
        findings_map.setdefault(f.clause_code, []).append(f)
    states_map = {s.clause_code: s for s in states_list}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "SimSun")

    title = doc.add_heading("合同审查报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    info_table = doc.add_table(rows=5, cols=2, style="Table Grid")
    info_data = [
        ("合同名称", contract.name),
        ("合同编号", contract.no or "-"),
        ("合同类型", ct.name if ct else (contract.type_code or "-")),
        ("审查立场", STANCE_LABEL.get(review.stance, review.stance) if review else "-"),
        ("审查状态", ver.status),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v

    doc.add_heading("审查结果汇总", level=1)
    high = sum(1 for f in findings_list if f.risk_level == "high")
    mid = sum(1 for f in findings_list if f.risk_level in ("medium", "mid"))
    low = sum(1 for f in findings_list if f.risk_level == "low")
    applied = sum(1 for s in states_list if s.applied)

    summary = doc.add_paragraph()
    summary.add_run(f"共发现 {len(findings_list)} 项风险（")
    run_h = summary.add_run(f"高风险 {high}")
    run_h.font.color.rgb = RGBColor(0xB3, 0x26, 0x1E)
    summary.add_run(f"、中风险 {mid}、低风险 {low}）。")
    summary.add_run(f"\n人工复核：已应用 AI 建议 {applied} 条。")

    doc.add_heading("逐条审查详情", level=1)
    for clause in clauses:
        cfs = findings_map.get(clause.code, [])
        st = states_map.get(clause.code)
        doc.add_heading(f"{clause.code} {clause.title}", level=2)
        p_text = doc.add_paragraph()
        p_text.add_run("条款内容：").bold = True
        p_text.add_run((clause.text or "")[:2000])
        if st and st.applied:
            p_app = doc.add_paragraph()
            p_app.add_run("已应用建议：").bold = True
            run_a = p_app.add_run("是（条款文本已按 AI 建议修改）")
            run_a.font.color.rgb = RGBColor(0x4A, 0x6B, 0x38)
        if st and st.note:
            p_note = doc.add_paragraph()
            p_note.add_run("人工批注：").bold = True
            p_note.add_run(st.note)
        for f in cfs:
            p_risk = doc.add_paragraph()
            p_risk.add_run(f"AI发现 [{RISK_LABEL.get(f.risk_level, f.risk_level)}]：").bold = True
            p_risk.add_run(f.finding)
            if f.suggestion:
                p_sug = doc.add_paragraph()
                p_sug.add_run("修改建议：").bold = True
                p_sug.add_run(f.suggestion)
            if f.legal_basis:
                refs = _format_legal_basis(f.legal_basis)
                if refs:
                    p_law = doc.add_paragraph()
                    p_law.add_run("法律依据：").bold = True
                    p_law.add_run(refs)
            if f.stance_note:
                p_st = doc.add_paragraph()
                p_st.add_run("立场分析：").bold = True
                p_st.add_run(f.stance_note)
        if not cfs:
            doc.add_paragraph("（无风险发现）")

    doc.add_paragraph("")
    doc.add_paragraph("— 报告由「明衡 AI 合同审阅系统」自动生成 —").alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------- docx → pdf（LibreOffice headless） ----------------

def _libreoffice_cmd() -> list[str]:
    """查找 libreoffice / soffice 可执行文件（Linux Docker、macOS 本地开发）。"""
    candidates = [
        "libreoffice",
        "soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
    ]
    for cmd in candidates:
        if cmd.startswith("/"):
            if Path(cmd).is_file():
                return [cmd]
        elif shutil.which(cmd):
            return [cmd]
    return []


def convert_to_pdf(docx_buf: io.BytesIO, out_dir: Path | None = None) -> io.BytesIO:
    """把 docx 字节流转成 pdf 字节流。依赖系统 libreoffice / soffice。"""
    lo_cmd = _libreoffice_cmd()
    if not lo_cmd:
        raise RuntimeError(
            "PDF 转换失败：服务器未安装 LibreOffice。"
            " macOS 本地开发可执行：brew install --cask libreoffice"
        )
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        src = td_path / "source.docx"
        with open(src, "wb") as f:
            f.write(docx_buf.getvalue())
        try:
            env = {**os.environ, "HOME": "/tmp"}
            subprocess.run(
                [*lo_cmd, "--headless", "--convert-to", "pdf", "--outdir", str(td_path), str(src)],
                check=True, capture_output=True, timeout=120, env=env,
            )
        except subprocess.CalledProcessError as e:
            err = (e.stderr or b"").decode(errors="replace")[-500:]
            raise RuntimeError(f"PDF 转换失败：LibreOffice 转换出错（{err.strip() or '未知错误'}）") from e
        except subprocess.TimeoutExpired as e:
            raise RuntimeError("PDF 转换失败：LibreOffice 转换超时") from e
        pdf_path = td_path / "source.pdf"
        if not pdf_path.exists():
            raise RuntimeError("PDF 转换失败：未生成输出文件")
        buf = io.BytesIO(pdf_path.read_bytes())
        buf.seek(0)
        return buf
