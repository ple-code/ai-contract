"""生成采购合同 Word 模板（含 docxtpl / Jinja2 占位符）。

运行：python -m app.seeds.seed_template
生成：app/api/app/templates/采购合同模板.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"
TEMPLATE_PATH = TEMPLATE_DIR / "采购合同模板.docx"


def _set_font(run, name: str, size: float, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _add_para(doc, text: str, font="宋体", size=12, bold=False,
              align=None, indent_chars=0, space_after=6, line_spacing=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if indent_chars:
        pf.first_line_indent = Pt(size * indent_chars)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        _set_font(run, font, size, bold)
    return p


def build_template():
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for s in doc.sections:
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(3.0)

    _add_para(doc, "采 购 合 同", font="黑体", size=22, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, line_spacing=1.5)

    _add_para(doc, "合同编号：{{ contract_no }}", font="仿宋", size=10.5,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, line_spacing=1.2)

    _add_para(doc, "签订地点：{{ sign_place }}        签订日期：{{ sign_date }}",
              font="仿宋", size=10.5, space_after=8, line_spacing=1.4)

    _add_para(doc, "甲方（采购方）：{{ party_a }}", font="宋体", size=12,
              space_after=4, line_spacing=1.5)
    _add_para(doc, "乙方（供应方）：{{ party_b }}", font="宋体", size=12,
              space_after=10, line_spacing=1.5)

    _add_para(doc, "根据《中华人民共和国民法典》及相关法律法规的规定，甲乙双方在平等、自愿、"
                   "公平和诚实信用的基础上，经友好协商，就货物采购事宜达成如下协议，以资共同遵守：",
              font="宋体", size=12, indent_chars=2, space_after=8, line_spacing=1.5)

    # 条款循环 —— {%p %} 段落级标签消除自身段落
    _add_para(doc, "{%p for c in clauses %}", space_after=0, line_spacing=1.0)
    _add_para(doc, "{{ c.code }}  {{ c.title }}", font="黑体", size=12, bold=True,
              space_after=2, line_spacing=1.5)
    _add_para(doc, "{{ c.text }}", font="宋体", size=12,
              indent_chars=2, space_after=6, line_spacing=1.5)
    _add_para(doc, "{%p endfor %}", space_after=0, line_spacing=1.0)

    _add_para(doc, "（以下无正文，为本合同签章页）", font="仿宋", size=10.5,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16, line_spacing=1.2)

    _add_para(doc, "甲方（盖章）：{{ party_a }}", font="宋体", size=12,
              space_after=10, line_spacing=1.5)
    _add_para(doc, "法定代表人 / 授权代表（签字）：____________________", font="宋体", size=12,
              space_after=10, line_spacing=1.5)
    _add_para(doc, "签订日期：________年______月______日", font="宋体", size=12,
              space_after=20, line_spacing=1.5)

    _add_para(doc, "乙方（盖章）：{{ party_b }}", font="宋体", size=12,
              space_after=10, line_spacing=1.5)
    _add_para(doc, "法定代表人 / 授权代表（签字）：____________________", font="宋体", size=12,
              space_after=10, line_spacing=1.5)
    _add_para(doc, "签订日期：________年______月______日", font="宋体", size=12,
              space_after=6, line_spacing=1.5)

    doc.save(TEMPLATE_PATH)
    return TEMPLATE_PATH


if __name__ == "__main__":
    p = build_template()
    print(f"已生成采购合同模板：{p}")
